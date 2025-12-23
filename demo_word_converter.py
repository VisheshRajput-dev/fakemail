#!/usr/bin/env python3
"""
Demo script to show Word document to HTML conversion
Run this to test the Word document conversion feature
"""

import os
import sys

# Add current directory to path so we can import the main script
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

try:
    from fake_email_sender_windows import FakeEmailSender
    print("[OK] Successfully imported FakeEmailSender")

    # Create a sample Word document for testing
    try:
        from docx import Document

        # Create a sample document
        doc = Document()
        doc.add_heading('Employee Confirmation Letter', 0)

        doc.add_paragraph('Dear Hiring Manager,')

        doc.add_paragraph('This is to confirm that the following employees have been selected:')

        # Create a table
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sr. No'
        hdr_cells[1].text = 'Employee Name'
        hdr_cells[2].text = 'Role'
        hdr_cells[3].text = 'Salary'

        # Add sample data
        rows = [
            ['1', 'John Doe', 'Developer', '$50,000'],
            ['2', 'Jane Smith', 'Designer', '$45,000'],
        ]

        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = cell_data

        doc.add_paragraph('We are pleased to have these employees on board.')

        # Save the document
        demo_file = 'demo_employee_letter.docx'
        doc.save(demo_file)
        print(f"[OK] Created demo Word document: {demo_file}")

    except ImportError:
        print("[INFO] python-docx not installed. Creating a placeholder message.")
        print("   To create demo documents, install with: pip install python-docx")

        # Create a simple demo HTML file instead
        demo_html = """
        <html>
        <body>
        <h1>Employee Confirmation Letter</h1>
        <p>Dear Hiring Manager,</p>
        <p>This is to confirm that the following employees have been selected:</p>
        <table border="1">
        <tr><th>Sr. No</th><th>Employee Name</th><th>Role</th><th>Salary</th></tr>
        <tr><td>1</td><td>John Doe</td><td>Developer</td><td>$50,000</td></tr>
        <tr><td>2</td><td>Jane Smith</td><td>Designer</td><td>$45,000</td></tr>
        </table>
        <p>We are pleased to have these employees on board.</p>
        </body>
        </html>
        """

        with open('demo_employee_letter.html', 'w') as f:
            f.write(demo_html)
        demo_file = 'demo_employee_letter.html'
        print(f"[OK] Created demo HTML file: {demo_file}")

    # Test the conversion
    sender = FakeEmailSender()

    if hasattr(sender, 'convert_word_to_html'):
        print("\n[TEST] Testing Word document conversion...")

        try:
            if demo_file.endswith('.docx'):
                html_content = sender.convert_word_to_html(demo_file)
                print("[OK] Word document conversion successful!")
                print("[PREVIEW] HTML Preview (first 200 characters):")
                print(html_content[:200] + "..." if len(html_content) > 200 else html_content)
            else:
                print("[INFO] Demo file is HTML, not testing conversion.")
                with open(demo_file, 'r') as f:
                    html_content = f.read()
                print("[OK] HTML content loaded successfully!")

        except Exception as e:
            print(f"[ERROR] Conversion error: {e}")
    else:
        print("[ERROR] convert_word_to_html method not found")

    print(f"\n[FILE] Demo file created: {demo_file}")
    print("[TARGET] You can now use this file with the email sender script!")

except ImportError as e:
    print(f"[ERROR] Import error: {e}")
    print("Make sure you're running this from the same directory as fake_email_sender_windows.py")
